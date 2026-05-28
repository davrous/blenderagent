"""Generate a feedback Word document summarizing our findings about the
Foundry Hosted Agent developer experience while building the
`fantasy-worlds-agent` Blender hosted agent.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


CODE_FONT = "Consolas"
BODY_FONT = "Calibri"

OUTPUT = Path(__file__).resolve().parent.parent / "Foundry-Hosted-Agent-Feedback-v2.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = BODY_FONT


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.bold = bold


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)


def add_code(doc: Document, code: str, language: str = "python") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    run.font.name = CODE_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def add_callout(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.name = BODY_FONT
    label_run.font.size = Pt(11)
    body_run = p.add_run(body)
    body_run.font.name = BODY_FONT
    body_run.font.size = Pt(11)


def build() -> Path:
    doc = Document()

    # --- Title -----------------------------------------------------------
    title = doc.add_heading("Foundry Hosted Agent – Developer Experience Feedback", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "Findings from building the `fantasy-worlds-agent` hosted agent (Python, "
        "Microsoft Agent Framework, Foundry Responses protocol) and getting "
        "observability/traces to appear in the Foundry portal."
    )
    sub_run.italic = True
    sub_run.font.name = BODY_FONT
    sub_run.font.size = Pt(11)

    # --- Context ---------------------------------------------------------
    add_heading(doc, "1. Context", level=1)
    add_paragraph(
        doc,
        "We are building a hosted agent (`kind: hosted`, name `fantasy-worlds-agent`) "
        "running as a container on Foundry. The agent drives a Blender server "
        "over a custom socket connection and exposes ~15 tools to the model. "
        "Deployment uses `agent.yaml` + Docker. The model is `gpt-5-mini`.",
    )
    add_paragraph(
        doc,
        "Goal of this iteration: get `gen_ai.*` traces (invoke_agent / chat / "
        "execute_tool) to appear in the Foundry portal **Traces** tab.",
    )

    # --- Problem 1 -------------------------------------------------------
    add_heading(doc, "2. Problem #1 — HTTP 400 \"Agent kind mismatch\"", level=1)
    add_paragraph(
        doc,
        "Our first implementation followed older samples and used:",
    )
    add_code(
        doc,
        "from agent_framework.azure import AzureAIClient\n"
        "from azure.ai.agentserver import from_agent_framework, InMemoryAgentSessionRepository\n"
        "\n"
        "client = AzureAIClient(\n"
        "    project_endpoint=..., model=..., credential=DefaultAzureCredential(),\n"
        ")\n"
        "agent = ChatAgent(client=client, name=\"fantasy-worlds-agent\", ...)\n"
        "server = from_agent_framework(agent, session_repository=InMemoryAgentSessionRepository())\n"
        "await server.run_async()\n",
    )
    add_paragraph(doc, "Symptom (in the running container):", bold=True)
    add_code(
        doc,
        'HTTP 400 from Foundry: "Agent kind mismatch. Expected hosted, got prompt."',
        language="text",
    )
    add_paragraph(doc, "Root cause:", bold=True)
    add_paragraph(
        doc,
        "`AzureAIClient._prepare_options` unconditionally calls "
        "`project_client.agents.create_version(...)` on every request. "
        "That registers the agent as `kind: prompt` in the Foundry project. "
        "But our `agent.yaml` declares `kind: hosted`. The runtime then refuses "
        "the request with the kind-mismatch error — no matter what value (or "
        "no value) we pass for `name=`.",
    )
    add_callout(
        doc,
        "Feedback for the team",
        "`AzureAIClient` is the wrong client for hosted agents but this is not "
        "obvious from its name or docs. The current samples README does not call "
        "this out clearly. Suggest either (a) renaming/scoping `AzureAIClient` to "
        "make it explicitly the prompt-agent client, or (b) detecting "
        "`kind: hosted` in agent.yaml and raising a clear local error before the "
        "round-trip.",
    )

    # --- Fix #1 ----------------------------------------------------------
    add_heading(doc, "3. Fix #1 — Migrate to the canonical hosted-agent stack", level=1)
    add_paragraph(
        doc,
        "We migrated to the stack used in "
        "`microsoft-foundry/foundry-samples/.../hosted-agents/agent-framework/responses/02-tools/main.py` "
        "and `.../08-observability/main.py`:",
    )
    add_code(
        doc,
        "from agent_framework import Agent\n"
        "from agent_framework.foundry import FoundryChatClient\n"
        "from agent_framework_foundry_hosting import ResponsesHostServer\n"
        "from azure.identity import DefaultAzureCredential   # SYNC\n"
        "\n"
        "client = FoundryChatClient(\n"
        "    project_endpoint=os.environ[\"FOUNDRY_PROJECT_ENDPOINT\"],\n"
        "    model=os.environ[\"AZURE_AI_MODEL_DEPLOYMENT_NAME\"],\n"
        "    credential=DefaultAzureCredential(),\n"
        ")\n"
        "agent = Agent(\n"
        "    client=client,\n"
        "    instructions=...,\n"
        "    tools=[...],\n"
        "    default_options={\"store\": False},\n"
        ")\n"
        "server = ResponsesHostServer(agent)\n"
        "await server.run_async()\n",
    )
    add_paragraph(doc, "Key differences from the AzureAIClient stack:", bold=True)
    add_bullet(doc, "`FoundryChatClient` never calls `agents.create_version` – it leaves the agent identity entirely to `agent.yaml`.")
    add_bullet(doc, "`Agent` (NOT `ChatAgent`) is what the installed `agent_framework` exports in 1.2.x.")
    add_bullet(doc, "No `name=` on `Agent(...)` — agent.yaml is the single source of truth for the hosted-agent name and kind.")
    add_bullet(doc, "`default_options={\"store\": False}` tells the service to manage conversation state.")
    add_bullet(doc, "`ResponsesHostServer(agent)` replaces `from_agent_framework(...)` + `InMemoryAgentSessionRepository()`.")

    # --- Problem 2 -------------------------------------------------------
    add_heading(doc, "4. Problem #2 — `ChatAgent` is not exported", level=1)
    add_paragraph(
        doc,
        "Multiple public samples and blog posts use `from agent_framework import ChatAgent`. "
        "In the version we installed (`agent-framework>=1.2.2`), `ChatAgent` is not exported "
        "from the top-level package — only `Agent` is.",
    )
    add_callout(
        doc,
        "Feedback for the team",
        "Either re-export `ChatAgent` as an alias of `Agent` (for back-compat with public "
        "samples), or do a sweep to retire the `ChatAgent` name from all current samples "
        "and READMEs. The mismatch costs every new developer at least one debugging cycle.",
    )

    # --- Problem 3 -------------------------------------------------------
    add_heading(
        doc,
        "5. Problem #3 — `FoundryChatClient` is not an async context manager",
        level=1,
    )
    add_paragraph(
        doc,
        "Based on the older `AzureAIClient` ergonomics (which IS an async context manager), "
        "we wrote:",
    )
    add_code(
        doc,
        "async with (\n"
        "    DefaultAzureCredential() as credential,\n"
        "    FoundryChatClient(\n"
        "        project_endpoint=PROJECT_ENDPOINT,\n"
        "        model=MODEL_DEPLOYMENT_NAME,\n"
        "        credential=credential,\n"
        "    ) as chat_client,\n"
        "):\n"
        "    agent = Agent(client=chat_client, ...)\n"
        "    server = ResponsesHostServer(agent)\n"
        "    await server.run_async()\n",
    )
    add_paragraph(doc, "Symptom:", bold=True)
    add_code(
        doc,
        "TypeError: 'FoundryChatClient' object does not support the asynchronous "
        "context manager protocol\n"
        "  File \"/app/main.py\", line 2046, in main\n"
        "    async with (",
        language="text",
    )
    add_paragraph(doc, "Root cause:", bold=True)
    add_paragraph(
        doc,
        "`FoundryChatClient` does not implement `__aenter__` / `__aexit__`. "
        "The canonical samples construct it plainly and rely on process lifetime "
        "for cleanup. The samples also use the SYNC `DefaultAzureCredential` from "
        "`azure.identity`, not the async one from `azure.identity.aio`.",
    )
    add_paragraph(doc, "Final working pattern:", bold=True)
    add_code(
        doc,
        "credential = DefaultAzureCredential()              # sync\n"
        "chat_client = FoundryChatClient(\n"
        "    project_endpoint=PROJECT_ENDPOINT,\n"
        "    model=MODEL_DEPLOYMENT_NAME,\n"
        "    credential=credential,\n"
        ")\n"
        "agent = Agent(client=chat_client, ...)\n"
        "server = ResponsesHostServer(agent)\n"
        "await server.run_async()\n",
    )
    add_callout(
        doc,
        "Feedback for the team",
        "Either implement `__aenter__`/`__aexit__` on `FoundryChatClient` (so the "
        "intuitive `async with` pattern works), or add a docstring / runtime "
        "`TypeError` message that points developers at the correct construction "
        "pattern. Today the type error is bare and gives no hint that the fix is "
        "\"just remove async with\".",
    )

    # --- Observability ---------------------------------------------------
    add_heading(doc, "6. Observability — getting traces into the Foundry portal", level=1)
    add_paragraph(
        doc,
        "Our previous code called `enable_instrumentation()` and "
        "`configure_azure_monitor()` in process. That works for self-hosted "
        "scenarios but is the WRONG model for hosted agents. The `08-observability` "
        "sample is env-var driven, with NO observability code:",
    )
    add_code(
        doc,
        "# agent.yaml\n"
        "env:\n"
        "  - name: ENABLE_INSTRUMENTATION\n"
        "    value: \"true\"\n"
        "  - name: ENABLE_SENSITIVE_DATA\n"
        "    value: \"true\"\n",
        language="yaml",
    )
    add_paragraph(doc, "Behavior in production:", bold=True)
    add_bullet(doc, "Foundry injects `APPLICATIONINSIGHTS_CONNECTION_STRING` into the container.")
    add_bullet(doc, "Foundry owns the OpenTelemetry TracerProvider for the process.")
    add_bullet(doc, "The agent framework emits OTel GenAI spans: `invoke_agent`, `chat`, `execute_tool`.")
    add_bullet(doc, "Traces appear in the Foundry portal **Traces** tab (next to Playground) — not just in App Insights.")
    add_callout(
        doc,
        "Feedback for the team",
        "This was the single biggest source of confusion. Public docs / blog posts "
        "still show `enable_instrumentation()` + `configure_azure_monitor()` for "
        "hosted agents, which silently double-instruments and can fight Foundry's "
        "TracerProvider. The hosted-agent docs should say, in one place and "
        "prominently: \"do NOT call `enable_instrumentation()` in a hosted agent — "
        "set `ENABLE_INSTRUMENTATION=true` in agent.yaml and let Foundry wire it.\"",
    )

    # --- Final layout ----------------------------------------------------
    add_heading(doc, "7. Final working layout", level=1)
    add_paragraph(doc, "main.py (relevant excerpt):", bold=True)
    add_code(
        doc,
        "from agent_framework import Agent\n"
        "from agent_framework.foundry import FoundryChatClient\n"
        "from agent_framework_foundry_hosting import ResponsesHostServer\n"
        "from azure.identity import DefaultAzureCredential\n"
        "\n"
        "async def main():\n"
        "    credential = DefaultAzureCredential()\n"
        "    chat_client = FoundryChatClient(\n"
        "        project_endpoint=PROJECT_ENDPOINT,\n"
        "        model=MODEL_DEPLOYMENT_NAME,\n"
        "        credential=credential,\n"
        "    )\n"
        "    agent = Agent(\n"
        "        client=chat_client,\n"
        "        middleware=[SceneIsolationMiddleware(ToolStatusMiddleware(), scene_manager)],\n"
        "        default_options={\"store\": False},\n"
        "        instructions=\"...\",\n"
        "        tools=[get_scene_info, create_object, ...],\n"
        "    )\n"
        "    server = ResponsesHostServer(agent)\n"
        "    await server.run_async()\n"
        "\n"
        "if __name__ == \"__main__\":\n"
        "    asyncio.run(main())\n",
    )
    add_paragraph(doc, "agent.yaml (relevant excerpt):", bold=True)
    add_code(
        doc,
        "kind: hosted\n"
        "name: fantasy-worlds-agent\n"
        "protocols:\n"
        "  - type: responses\n"
        "    version: 1.0.0\n"
        "env:\n"
        "  - name: PROJECT_ENDPOINT\n"
        "    value: \"https://...\"\n"
        "  - name: MODEL_DEPLOYMENT_NAME\n"
        "    value: \"gpt-5-mini\"\n"
        "  - name: ENABLE_INSTRUMENTATION\n"
        "    value: \"true\"\n"
        "  - name: ENABLE_SENSITIVE_DATA\n"
        "    value: \"true\"\n",
        language="yaml",
    )
    add_paragraph(doc, "requirements.txt (relevant):", bold=True)
    add_code(
        doc,
        "agent-framework>=1.2.2\n"
        "agent-framework-foundry-hosting\n"
        "azure-identity>=1.25.3\n",
        language="text",
    )

    # --- Top recommendations --------------------------------------------
    add_heading(doc, "8. Top recommendations for the team", level=1)
    add_bullet(
        doc,
        "Make the hosted-agent vs prompt-agent distinction explicit in the SDK. "
        "Today they share a name (`agent_framework`) and similar APIs but only "
        "one of them works for `kind: hosted`. A clear local validation against "
        "agent.yaml would catch this before deploy.",
    )
    add_bullet(
        doc,
        "Publish ONE canonical \"hosted agent with tools + observability\" sample "
        "(today it's split across 02-tools and 08-observability) and link it from "
        "the top of the Foundry hosted-agent docs.",
    )
    add_bullet(
        doc,
        "Either support `async with FoundryChatClient(...)` or document/error "
        "clearly that it's not supported.",
    )
    add_bullet(
        doc,
        "Re-export `ChatAgent` or update all docs/samples/blog posts to use `Agent`.",
    )
    add_bullet(
        doc,
        "Make it crystal clear in the hosted-agent docs that observability is "
        "env-var driven (`ENABLE_INSTRUMENTATION=true`) and that calling "
        "`enable_instrumentation()` in process is incorrect for hosted agents.",
    )
    add_bullet(
        doc,
        "Improve the kind-mismatch error from Foundry to include the offending "
        "agent name and the kinds compared (\"expected hosted because agent.yaml "
        "declares it; got prompt because AzureAIClient registered a new version\").",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")

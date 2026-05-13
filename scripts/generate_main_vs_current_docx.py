"""Generate a Word document comparing the `main` branch implementation of
`fantasy-worlds-agent` against the current working-tree implementation,
focused on the **SDK choice** and **telemetry/observability setup**.

The intended audience is the Foundry Hosted Agent Developer Experience team.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


CODE_FONT = "Consolas"
BODY_FONT = "Calibri"

OUTPUT = (
    Path(__file__).resolve().parent.parent
    / "Foundry-Hosted-Agent-Main-vs-Current.docx"
)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = BODY_FONT


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code)
    run.font.name = CODE_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def add_callout(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.name = BODY_FONT
    label_run.font.size = Pt(11)
    body_run = p.add_run(body)
    body_run.font.name = BODY_FONT
    body_run.font.size = Pt(11)


def add_two_col_table(doc: Document, header_left: str, header_right: str, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for cell, text in ((hdr[0], header_left), (hdr[1], header_right)):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(11)
    for i, (left, right) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for cell, text in ((cells[0], left), (cells[1], right)):
            run = cell.paragraphs[0].add_run(text)
            run.font.name = BODY_FONT
            run.font.size = Pt(10)


def build() -> Path:
    doc = Document()

    # --- Title ----------------------------------------------------------
    title = doc.add_heading(
        "Foundry Hosted Agent SDK & Telemetry — `main` vs current working tree",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_paragraph(
        doc,
        "Project: `fantasy-worlds-agent` (Python hosted agent, Foundry Responses "
        "protocol, gpt-5-mini, custom Blender backend). This document summarizes "
        "the evolution of the agent host code between the version currently "
        "committed to the `main` branch and the version we have in the working "
        "tree today, focusing on (1) the Microsoft Agent Framework SDK surface "
        "we depend on, and (2) the OpenTelemetry / Application Insights "
        "telemetry setup.",
        italic=True,
    )

    # --- TL;DR -----------------------------------------------------------
    add_heading(doc, "1. TL;DR", 1)
    add_bullet(
        doc,
        "We migrated off `azure.ai.agentserver.agentframework` + "
        "`agent_framework.azure.AzureAIClient` and onto the new canonical "
        "hosted-agent stack: `agent_framework.foundry.FoundryChatClient` + "
        "`agent_framework_foundry_hosting.ResponsesHostServer`.",
    )
    add_bullet(
        doc,
        "Telemetry went from \"no setup at all\" (no spans appeared in Foundry "
        "Traces) to the env-var driven model documented in the 08-observability "
        "sample (`ENABLE_INSTRUMENTATION=true` in `agent.yaml`), plus one "
        "explicit `enable_instrumentation()` call at process start.",
    )
    add_bullet(
        doc,
        "We had to discover several non-obvious behaviour differences along "
        "the way (kind-mismatch HTTP 400, `ChatAgent` not exported, "
        "`FoundryChatClient` not being an async context manager, "
        "`ResponsesHostServer` not auto-enabling the framework tracer). They "
        "are detailed in section 4.",
    )

    # --- Side-by-side ---------------------------------------------------
    add_heading(doc, "2. Side-by-side summary", 1)
    add_two_col_table(
        doc,
        "`main` branch",
        "Current working tree",
        [
            (
                "agent-framework + azure-ai-agentserver-agentframework==1.0.0b17",
                "agent-framework>=1.2.2 + agent-framework-foundry-hosting",
            ),
            (
                "from agent_framework.azure import AzureAIClient",
                "from agent_framework.foundry import FoundryChatClient",
            ),
            (
                "from azure.ai.agentserver.agentframework import from_agent_framework",
                "from agent_framework_foundry_hosting import ResponsesHostServer",
            ),
            (
                "from azure.identity.aio import DefaultAzureCredential  (async)",
                "from azure.identity import DefaultAzureCredential  (sync)",
            ),
            (
                "Async context manager: `async with AzureAIClient(...).as_agent(name=..., tools=...) as agent:`",
                "Plain construction: `Agent(client=chat_client, tools=..., default_options={\"store\": False})` — no `async with`",
            ),
            (
                "Agent name passed in code: `as_agent(name=\"BlenderSceneAgent\", ...)`",
                "Agent name lives ONLY in agent.yaml (`name: fantasy-worlds-agent`); no `name=` in code",
            ),
            (
                "`server = from_agent_framework(agent, session_repository=InMemoryAgentSessionRepository())`",
                "`server = ResponsesHostServer(agent)`",
            ),
            (
                "No observability setup",
                "`enable_instrumentation(enable_sensitive_data=…)` called at startup",
            ),
            (
                "agent.yaml: no observability env vars",
                "agent.yaml adds `ENABLE_INSTRUMENTATION=true` and `ENABLE_SENSITIVE_DATA=true`",
            ),
            (
                "Result: no `gen_ai.*` spans visible in Foundry Traces tab",
                "Result: `invoke_agent`, `chat`, `execute_tool` spans visible in Foundry Traces and App Insights",
            ),
        ],
    )

    # --- main.py before/after -------------------------------------------
    add_heading(doc, "3. main.py — relevant excerpts", 1)

    add_paragraph(doc, "3.1 `main` branch (before):", bold=True)
    add_code(
        doc,
        "from agent_framework.azure import AzureAIClient\n"
        "from azure.ai.agentserver.agentframework import from_agent_framework\n"
        "from azure.ai.agentserver.agentframework.persistence import (\n"
        "    InMemoryAgentSessionRepository,\n"
        ")\n"
        "from azure.identity.aio import DefaultAzureCredential\n"
        "\n"
        "async def main():\n"
        "    async with (\n"
        "        DefaultAzureCredential() as credential,\n"
        "        AzureAIClient(\n"
        "            project_endpoint=PROJECT_ENDPOINT,\n"
        "            model_deployment_name=MODEL_DEPLOYMENT_NAME,\n"
        "            credential=credential,\n"
        "        ).as_agent(\n"
        "            name=\"BlenderSceneAgent\",\n"
        "            middleware=[SceneIsolationMiddleware(ToolStatusMiddleware(), scene_manager)],\n"
        "            instructions=\"...\",\n"
        "            tools=[get_scene_info, create_object, ...],\n"
        "        ) as agent,\n"
        "    ):\n"
        "        server = from_agent_framework(\n"
        "            agent, session_repository=InMemoryAgentSessionRepository()\n"
        "        )\n"
        "        await server.run_async()\n",
    )

    add_paragraph(doc, "3.2 Working tree (current):", bold=True)
    add_code(
        doc,
        "from agent_framework import Agent\n"
        "from agent_framework.foundry import FoundryChatClient\n"
        "from agent_framework.observability import enable_instrumentation\n"
        "from agent_framework_foundry_hosting import ResponsesHostServer\n"
        "from azure.identity import DefaultAzureCredential  # SYNC\n"
        "\n"
        "async def main():\n"
        "    # Telemetry: opt-in to GenAI span instrumentation. The hosting\n"
        "    # runtime configures the Azure Monitor exporter from\n"
        "    # APPLICATIONINSIGHTS_CONNECTION_STRING but does NOT auto-enable\n"
        "    # the framework tracer/meter, unlike the older\n"
        "    # `from_agent_framework` host.\n"
        "    try:\n"
        "        sensitive = os.getenv(\"ENABLE_SENSITIVE_DATA\", \"\").lower() in (\"1\", \"true\", \"yes\")\n"
        "        enable_instrumentation(enable_sensitive_data=sensitive or None)\n"
        "    except Exception as obs_exc:\n"
        "        logger.warning(\"enable_instrumentation() failed: %s\", obs_exc)\n"
        "\n"
        "    credential = DefaultAzureCredential()\n"
        "    chat_client = FoundryChatClient(\n"
        "        project_endpoint=PROJECT_ENDPOINT,\n"
        "        model=MODEL_DEPLOYMENT_NAME,\n"
        "        credential=credential,\n"
        "    )\n"
        "    agent = Agent(\n"
        "        client=chat_client,\n"
        "        middleware=[SceneIsolationMiddleware(ToolStatusMiddleware(), scene_manager)],\n"
        "        default_options={\"store\": False},\n"
        "        instructions=\"...\",\n"
        "        tools=[get_scene_info, create_object, ...],\n"
        "    )\n"
        "    server = ResponsesHostServer(agent)\n"
        "    await server.run_async()\n",
    )

    # --- Frictions ------------------------------------------------------
    add_heading(doc, "4. Frictions we hit while migrating", 1)

    add_heading(doc, "4.1 HTTP 400 \"Agent kind mismatch\" with AzureAIClient", 2)
    add_paragraph(
        doc,
        "`AzureAIClient._prepare_options` calls "
        "`project_client.agents.create_version(...)` on every request, "
        "registering the agent as `kind: prompt` in the Foundry project. "
        "Our `agent.yaml` declares `kind: hosted`, so the runtime rejected "
        "every request with HTTP 400 regardless of the `name=` value passed.",
    )
    add_callout(
        doc,
        "Suggestion",
        "Either (a) clearly label `AzureAIClient` as the prompt-agent client, "
        "or (b) detect `kind: hosted` from agent.yaml and emit a local, actionable "
        "error before the first HTTP round-trip. The current error gives no hint "
        "about the prompt-vs-hosted distinction.",
    )

    add_heading(doc, "4.2 `ChatAgent` is not exported from `agent_framework`", 2)
    add_paragraph(
        doc,
        "Public samples and blog posts widely use `from agent_framework import ChatAgent`. "
        "In `agent-framework>=1.2.2` only `Agent` is exported. Several iterations were "
        "spent chasing this `ImportError`.",
    )
    add_callout(
        doc,
        "Suggestion",
        "Re-export `ChatAgent` as an alias of `Agent`, or remove `ChatAgent` from "
        "all current samples and docs. The mismatch costs every new developer at "
        "least one debugging cycle.",
    )

    add_heading(doc, "4.3 `FoundryChatClient` is not an async context manager", 2)
    add_paragraph(
        doc,
        "Coming from `AzureAIClient` (which is an async context manager), "
        "`async with FoundryChatClient(...) as client:` raises a bare `TypeError: "
        "'FoundryChatClient' object does not support the asynchronous context "
        "manager protocol`. The canonical samples construct it plainly and rely "
        "on process lifetime for cleanup, with the SYNC `DefaultAzureCredential`.",
    )
    add_callout(
        doc,
        "Suggestion",
        "Implement `__aenter__` / `__aexit__` on `FoundryChatClient` for parity, "
        "or surface a more actionable error pointing developers at the documented "
        "construction pattern.",
    )

    add_heading(doc, "4.4 `ResponsesHostServer` does NOT auto-enable the framework tracer", 2)
    add_paragraph(
        doc,
        "The previous host, `azure.ai.agentserver.agentframework.from_agent_framework`, "
        "implicitly enabled the framework's GenAI instrumentation. "
        "`ResponsesHostServer` configures the Azure Monitor exporter (you can see "
        "\"Application Insights trace exporter configured.\" in the container logs) "
        "but does NOT enable the framework tracer/meter. Without an explicit "
        "`enable_instrumentation()` call, the Foundry portal shows host-level "
        "spans only — `invoke_agent`, `chat` and `execute_tool` spans never appear.",
    )
    add_callout(
        doc,
        "Suggestion",
        "Either have `ResponsesHostServer` auto-call `enable_instrumentation()` "
        "when `ENABLE_INSTRUMENTATION=true` (so the env var is sufficient, "
        "matching the 08-observability sample's intent), or document very "
        "prominently that hosted agents must call `enable_instrumentation()` "
        "explicitly at startup. Today the env var alone is misleading — it "
        "switches on the exporter but not the spans.",
    )

    add_heading(doc, "4.5 Stream-finalization `ValueError: Token … was created in a different Context`", 2)
    add_paragraph(
        doc,
        "After enabling instrumentation, every streamed turn ended with a "
        "`ValueError: <Token … was created in a different Context>` raised from "
        "`agent_framework.observability._finalize_stream`. The cleanup hook does "
        "`Token.reset()` on a context variable; `Token.reset()` checks Context "
        "*identity*, and `asyncio.create_task` always produces a brand-new "
        "Context for the child task — even when `context=copy_context()` is "
        "passed. The exception fires AFTER all stream data has been delivered, "
        "but our middleware was forwarding it to the consumer, producing a "
        "spurious \"transient error\" banner at the end of every turn. We added "
        "a targeted `except ValueError` to swallow it.",
    )
    add_callout(
        doc,
        "Suggestion",
        "In `_finalize_stream`, guard the `Token.reset()` with a try/except and "
        "log at debug — this is purely a telemetry cleanup artifact and should "
        "never reach end users. As-is, any consumer that wraps the framework's "
        "stream in a background task will see this error.",
    )

    # --- Telemetry deep-dive --------------------------------------------
    add_heading(doc, "5. Telemetry / observability — what changed", 1)
    add_paragraph(doc, "5.1 `agent.yaml` additions:", bold=True)
    add_code(
        doc,
        "environment_variables:\n"
        "  # ... existing vars ...\n"
        "  - name: ENABLE_INSTRUMENTATION\n"
        "    value: \"true\"\n"
        "  - name: ENABLE_SENSITIVE_DATA\n"
        "    value: \"true\"\n",
    )
    add_paragraph(
        doc,
        "Foundry injects `APPLICATIONINSIGHTS_CONNECTION_STRING` into the "
        "container at runtime, and `ResponsesHostServer` wires the Azure Monitor "
        "trace exporter using it. `ENABLE_SENSITIVE_DATA=true` causes the "
        "framework to attach prompt/completion text to `gen_ai.*` spans, which "
        "powers the Foundry portal \"Messages\" view.",
    )
    add_paragraph(doc, "5.2 main.py additions:", bold=True)
    add_code(
        doc,
        "from agent_framework.observability import enable_instrumentation\n"
        "\n"
        "sensitive = os.getenv(\"ENABLE_SENSITIVE_DATA\", \"\").lower() in (\"1\", \"true\", \"yes\")\n"
        "enable_instrumentation(enable_sensitive_data=sensitive or None)\n",
    )
    add_paragraph(doc, "5.3 What we did NOT need (despite older docs/blog posts saying so):", bold=True)
    add_bullet(doc, "No `configure_azure_monitor(...)` call — Foundry/ResponsesHostServer handles the exporter side.")
    add_bullet(doc, "No manual `TracerProvider` or OTLP exporter setup.")
    add_bullet(doc, "No `logging.basicConfig` integration with OTel logs.")
    add_callout(
        doc,
        "Feedback",
        "The intersection of the env-var model (agent.yaml flags), the explicit "
        "`enable_instrumentation()` requirement, and the hands-off exporter "
        "wiring is not documented in one place. Today this knowledge is split "
        "between 02-tools/main.py, 08-observability/agent.yaml and the framework "
        "source. A single \"hosted agent with tools + telemetry\" sample, "
        "explicitly contrasted with the self-hosted/standalone scenarios, would "
        "save every team this entire migration.",
    )

    # --- Dependencies ---------------------------------------------------
    add_heading(doc, "6. requirements.txt — full diff", 1)
    add_paragraph(doc, "Before (`main`):", bold=True)
    add_code(
        doc,
        "azure-ai-agentserver-agentframework==1.0.0b17\n"
        "# azure-ai-projects pinned transitively by agent-framework-core==1.0.0rc3 to 2.0.0b4\n"
        "python-dotenv>=1.2.2\n"
        "azure-identity>=1.25.3\n"
        "azure-storage-blob>=12.28.0\n",
    )
    add_paragraph(doc, "After (current):", bold=True)
    add_code(
        doc,
        "agent-framework>=1.2.2\n"
        "agent-framework-foundry-hosting\n"
        "python-dotenv>=1.2.2\n"
        "azure-identity>=1.25.3\n"
        "azure-storage-blob>=12.28.0\n",
    )
    add_paragraph(
        doc,
        "Note: `agent-framework-foundry-hosting` is not pinned because we have "
        "not yet identified a stable release line; we are currently picking up "
        "whatever the index serves. A pinned version range from the team would "
        "be appreciated.",
    )

    # --- Top recommendations --------------------------------------------
    add_heading(doc, "7. Top asks for the DX team", 1)
    add_bullet(
        doc,
        "Publish ONE canonical \"hosted agent with tools + observability\" "
        "sample that covers the full picture (we had to merge 02-tools and "
        "08-observability mentally).",
    )
    add_bullet(
        doc,
        "Either make `ResponsesHostServer` auto-enable the framework tracer "
        "when `ENABLE_INSTRUMENTATION=true`, or document the explicit "
        "`enable_instrumentation()` requirement very prominently. Today the "
        "env var alone is necessary but not sufficient.",
    )
    add_bullet(
        doc,
        "Pin or version-range `agent-framework-foundry-hosting` in the public "
        "samples so downstream teams have a known-good combination.",
    )
    add_bullet(
        doc,
        "Fix or document the `Token … created in a different Context` "
        "stream-finalization `ValueError`. It produces spurious end-of-turn "
        "error banners for anyone using `asyncio.create_task` to bridge the "
        "framework stream into their own transport.",
    )
    add_bullet(
        doc,
        "Detect and reject `kind: hosted` agent.yaml + `AzureAIClient` "
        "combination locally with an actionable error. Today it round-trips "
        "and dies with a generic HTTP 400 \"Agent kind mismatch\".",
    )
    add_bullet(
        doc,
        "Re-export `ChatAgent` from `agent_framework` (alias of `Agent`) or "
        "purge it from samples/docs.",
    )
    add_bullet(
        doc,
        "Add `__aenter__`/`__aexit__` to `FoundryChatClient` to match the "
        "`AzureAIClient` ergonomics that the rest of the Azure SDKs use.",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
